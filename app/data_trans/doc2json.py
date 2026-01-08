"""
Word 文档解析器
支持 .docx 格式，.doc 格式需要转换
"""
from typing import Union, Dict, Any, List
from pathlib import Path

from .base_parser import BaseParser


class DocParser(BaseParser):
    """Word 文档解析器"""
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    
    def __init__(self):
        self._Document = None
    
    def _get_docx_lib(self):
        """懒加载 python-docx"""
        if self._Document is None:
            try:
                from docx import Document
                self._Document = Document
            except ImportError:
                raise ImportError("请安装 python-docx: pip install python-docx")
        return self._Document
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        解析 Word 文档
        """
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        
        try:
            if ext == '.docx':
                content = self._parse_docx(file_path)
            elif ext == '.doc':
                # .doc 是旧格式，需要转换或使用其他库
                content = self._parse_doc(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {ext}")
            
            return {
                "success": True,
                "file_path": str(file_path),
                "content": content,
                "error": None
            }
            
        except Exception as e:
            return {
                "success": False,
                "file_path": str(file_path),
                "content": None,
                "error": str(e)
            }
    
    def _parse_docx(self, file_path: Path) -> str:
        """解析 .docx 文件"""
        Document = self._get_docx_lib()
        doc = Document(file_path)
        
        paragraphs = []
        
        # 提取正文段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # 提取表格内容 (施教区文件经常有表格)
        for table in doc.tables:
            table_text = self._extract_table(table)
            if table_text:
                paragraphs.append(table_text)
        
        return "\n".join(paragraphs)
    
    def _extract_table(self, table) -> str:
        """提取表格内容为文本"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # 去重（合并单元格会导致重复）
            unique_cells = list(dict.fromkeys(cells))
            rows.append(" | ".join(unique_cells))
        return "\n".join(rows)
    
    def _parse_doc(self, file_path: Path) -> str:
        """
        解析 .doc 文件 (旧版 Word 格式)
        
        方案1: 使用 antiword (需要系统安装)
        方案2: 使用 textract
        方案3: 使用 LibreOffice 转换为 docx
        """
        # 尝试使用 textract
        try:
            import textract
            text = textract.process(str(file_path)).decode('utf-8')
            return text
        except ImportError:
            pass
        
        # 尝试使用 subprocess 调用 antiword
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        raise ImportError(
            ".doc 格式需要额外依赖。请选择以下方案之一：\n"
            "1. 安装 textract: pip install textract\n"
            "2. 安装 antiword: apt-get install antiword\n"
            "3. 将 .doc 文件转换为 .docx 格式"
        )
